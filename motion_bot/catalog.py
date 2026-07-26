"""Template catalog: register imported and sample motion templates."""

from __future__ import annotations

import re
import shutil
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import yaml

from motion_bot.paths import (
    LIBRARY_TEMPLATES_DIR,
    MANIFEST_PATH,
    SAMPLE_TEMPLATES_DIR,
    TEMPLATES_DIR,
    ensure_runtime_dirs,
)

# Capture Jinja names from {{ var }}, {{ var|filter }}, and {% for x in items %}
PLACEHOLDER_RE = re.compile(
    r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_\.]*)"
    r"|\{%\s*(?:if|elif)\s+([a-zA-Z_][a-zA-Z0-9_\.]*)"
    r"|\{%\s*for\s+[a-zA-Z_][a-zA-Z0-9_]*\s+in\s+([a-zA-Z_][a-zA-Z0-9_\.]*)"
)


@dataclass
class TemplateEntry:
    id: str
    name: str
    path: Path
    source: str  # "library" | "sample" | "custom"
    description: str = ""
    jurisdiction: str = ""
    motion_type: str = ""
    placeholders: list[str] = field(default_factory=list)
    notes: str = ""

    def to_dict(self) -> dict[str, Any]:
        return {
            "id": self.id,
            "name": self.name,
            "path": str(self.path.relative_to(TEMPLATES_DIR))
            if self.path.is_relative_to(TEMPLATES_DIR)
            else str(self.path),
            "source": self.source,
            "description": self.description,
            "jurisdiction": self.jurisdiction,
            "motion_type": self.motion_type,
            "placeholders": self.placeholders,
            "notes": self.notes,
        }


def _slugify(value: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", value.strip().lower()).strip("-")
    return slug or "template"


def _normalize_source(source: str) -> str:
    return source if source else "custom"


def _load_manifest() -> dict[str, Any]:
    ensure_runtime_dirs()
    if not MANIFEST_PATH.exists():
        return {"templates": []}
    with MANIFEST_PATH.open(encoding="utf-8") as fh:
        data = yaml.safe_load(fh) or {}
    if "templates" not in data or not isinstance(data["templates"], list):
        data["templates"] = []
    return data


def _save_manifest(data: dict[str, Any]) -> None:
    ensure_runtime_dirs()
    with MANIFEST_PATH.open("w", encoding="utf-8") as fh:
        yaml.safe_dump(data, fh, sort_keys=False, allow_unicode=True)


def discover_placeholders(docx_path: Path) -> list[str]:
    """Scan a .docx for Jinja-style placeholders used by docxtpl."""
    from docx import Document

    doc = Document(str(docx_path))
    chunks: list[str] = []
    for para in doc.paragraphs:
        chunks.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    for section in doc.sections:
        for part in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
        ):
            if part is None:
                continue
            for para in part.paragraphs:
                chunks.append(para.text)

    found: set[str] = set()
    for text in chunks:
        for match in PLACEHOLDER_RE.finditer(text):
            name = next((g for g in match.groups() if g), None)
            if name:
                found.add(name.split(".")[0])
    return sorted(found)


def _resolve_path(rel_or_abs: str | Path) -> Path:
    path = Path(rel_or_abs)
    if path.is_absolute():
        return path
    candidate = TEMPLATES_DIR / path
    if candidate.exists():
        return candidate
    return (TEMPLATES_DIR.parent / path).resolve()


def list_templates() -> list[TemplateEntry]:
    """Return registered templates; auto-register sample/library .docx files if missing."""
    ensure_runtime_dirs()
    data = _load_manifest()
    by_id: dict[str, dict[str, Any]] = {
        t["id"]: t for t in data["templates"] if isinstance(t, dict) and "id" in t
    }

    # Auto-discover unregistered .docx under sample/ and library/
    for source, directory in (
        ("sample", SAMPLE_TEMPLATES_DIR),
        ("library", LIBRARY_TEMPLATES_DIR),
    ):
        if not directory.exists():
            continue
        for docx in sorted(directory.glob("*.docx")):
            if docx.name.startswith("~$"):
                continue
            tid = _slugify(docx.stem)
            if tid in by_id:
                continue
            placeholders = discover_placeholders(docx)
            entry = {
                "id": tid,
                "name": docx.stem.replace("_", " ").replace("-", " ").title(),
                "path": str(docx.relative_to(TEMPLATES_DIR)),
                "source": source,
                "description": f"Auto-discovered {source} template",
                "jurisdiction": "",
                "motion_type": "",
                "placeholders": placeholders,
                "notes": "",
            }
            data["templates"].append(entry)
            by_id[tid] = entry

    entries: list[TemplateEntry] = []
    for raw in data["templates"]:
        path = _resolve_path(raw["path"])
        entries.append(
            TemplateEntry(
                id=raw["id"],
                name=raw.get("name", raw["id"]),
                path=path,
                source=_normalize_source(raw.get("source", "custom")),
                description=raw.get("description", ""),
                jurisdiction=raw.get("jurisdiction", ""),
                motion_type=raw.get("motion_type", ""),
                placeholders=list(raw.get("placeholders") or []),
                notes=raw.get("notes", ""),
            )
        )
    return entries


def get_template(template_id: str) -> TemplateEntry:
    for entry in list_templates():
        if entry.id == template_id:
            if not entry.path.exists():
                raise FileNotFoundError(
                    f"Template '{template_id}' is registered but file is missing: {entry.path}"
                )
            return entry
    known = ", ".join(e.id for e in list_templates()) or "(none)"
    raise KeyError(f"Unknown template_id '{template_id}'. Known: {known}")


def import_template(
    source_path: Path,
    *,
    template_id: str | None = None,
    name: str | None = None,
    description: str = "",
    jurisdiction: str = "",
    motion_type: str = "",
    notes: str = (
        "Imported motion template. Place Jinja placeholders "
        "like {{ case_number }} before generating."
    ),
) -> TemplateEntry:
    """Copy a user .docx into templates/library and register it."""
    ensure_runtime_dirs()
    source_path = Path(source_path).expanduser().resolve()
    if not source_path.exists():
        raise FileNotFoundError(f"No such file: {source_path}")
    if source_path.suffix.lower() != ".docx":
        raise ValueError("Only .docx templates are supported. Convert .doc files first.")

    tid = template_id or _slugify(source_path.stem)
    dest = LIBRARY_TEMPLATES_DIR / f"{tid}.docx"
    if source_path.resolve() != dest.resolve():
        shutil.copy2(source_path, dest)

    placeholders = discover_placeholders(dest)
    data = _load_manifest()
    data["templates"] = [t for t in data["templates"] if t.get("id") != tid]
    entry = {
        "id": tid,
        "name": name or source_path.stem.replace("_", " ").title(),
        "path": str(dest.relative_to(TEMPLATES_DIR)),
        "source": "library",
        "description": description or "Imported motion template",
        "jurisdiction": jurisdiction,
        "motion_type": motion_type,
        "placeholders": placeholders,
        "notes": notes,
    }
    data["templates"].append(entry)
    _save_manifest(data)
    return get_template(tid)


def refresh_placeholders(template_id: str) -> TemplateEntry:
    """Re-scan a template for placeholders and update the manifest."""
    entry = get_template(template_id)
    placeholders = discover_placeholders(entry.path)
    data = _load_manifest()
    for raw in data["templates"]:
        if raw.get("id") == template_id:
            raw["placeholders"] = placeholders
            break
    _save_manifest(data)
    return get_template(template_id)
