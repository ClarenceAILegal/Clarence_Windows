#!/usr/bin/env python3
"""Create a sample court-motion Word template with Jinja placeholders."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "templates" / "sample" / "sample-motion-to-compel.docx"


def set_run_font(run, *, size=12, bold=False, font="Times New Roman") -> None:
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold


def add_centered(doc: Document, text: str, *, bold=False, size=12, space_after=6) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_body(doc: Document, text: str, *, bold=False, first_line_indent=True) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(0)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    run = p.add_run(text)
    set_run_font(run, bold=bold)


def main() -> int:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    add_centered(doc, "{{ court_name }}", bold=True, size=12)
    add_centered(doc, "{{ court_division }}", bold=False, size=12)
    add_centered(doc, "{{ county }} County, {{ state }}", bold=False, size=12, space_after=18)

    # Caption block
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run("{{ plaintiff }},")
    set_run_font(run, bold=True)
    p = doc.add_paragraph()
    run = p.add_run("          Plaintiff,")
    set_run_font(run)
    p = doc.add_paragraph()
    run = p.add_run("v.")
    set_run_font(run)
    p = doc.add_paragraph()
    run = p.add_run("{{ defendant }},")
    set_run_font(run, bold=True)
    p = doc.add_paragraph()
    run = p.add_run("          Defendant.")
    set_run_font(run)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Case No. {{ case_number }}")
    set_run_font(run, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Judge: {{ judge }}")
    set_run_font(run)

    add_centered(doc, "{{ motion_title|upper }}", bold=True, size=12, space_after=18)

    add_body(
        doc,
        "COMES NOW {{ movant_name }} (\"Movant\"), by and through undersigned counsel, "
        "and respectfully moves this Court for {{ relief_sought }}, and in support thereof states as follows:",
        first_line_indent=False,
    )

    add_body(doc, "I. INTRODUCTION", bold=True, first_line_indent=False)
    add_body(
        doc,
        "This Motion seeks relief on the grounds set forth below. Movant is {{ movant_role }} "
        "in the above-captioned action.",
    )

    add_body(doc, "II. FACTUAL BACKGROUND", bold=True, first_line_indent=False)
    add_body(doc, "{{ factual_background }}")

    add_body(doc, "III. LEGAL ARGUMENT", bold=True, first_line_indent=False)
    add_body(doc, "{{ legal_argument }}")

    add_body(doc, "IV. PRAYER FOR RELIEF", bold=True, first_line_indent=False)
    add_body(doc, "{{ prayer_for_relief }}")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run("Respectfully submitted this {{ filing_date }},")
    set_run_font(run)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("{{ counsel_firm }}")
    set_run_font(run, bold=True)

    for line in (
        "{{ counsel_name }}",
        "Bar No. {{ counsel_bar }}",
        "{{ counsel_address }}",
        "Tel: {{ counsel_phone }}",
        "Email: {{ counsel_email }}",
        "Counsel for {{ movant_name }}",
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(line)
        set_run_font(run)

    add_body(doc, "CERTIFICATE OF SERVICE", bold=True, first_line_indent=False)
    add_body(doc, "{{ certificate_of_service }}")

    add_body(doc, "HEARING", bold=True, first_line_indent=False)
    add_body(
        doc,
        "Hearing requested: {{ hearing_date }} at {{ hearing_time }}, {{ hearing_location }}.",
        first_line_indent=False,
    )

    add_body(doc, "EXHIBITS", bold=True, first_line_indent=False)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run(
        "{% for exhibit in exhibits %}{{ exhibit }}\n{% endfor %}"
    )
    set_run_font(run)

    doc.save(str(OUT))
    print(f"Wrote sample template: {OUT}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
