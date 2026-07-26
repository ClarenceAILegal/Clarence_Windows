#!/usr/bin/env python3
"""Build fillable Word templates from Lexis FL eyewitness-ID motion forms 8.08 / 8.09."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "templates" / "lexis"


def _font(run, *, size=12, bold=False, italic=False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def _p(
    doc: Document,
    text: str = "",
    *,
    bold=False,
    italic=False,
    center=False,
    justify=False,
    space_before=0,
    space_after=6,
    first_indent=False,
    double=True,
) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = (
        WD_LINE_SPACING.DOUBLE if double else WD_LINE_SPACING.SINGLE
    )
    if first_indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    if text:
        run = p.add_run(text)
        _font(run, bold=bold, italic=italic)


def _base_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    return doc


def _caption(doc: Document) -> None:
    _p(doc, "IN THE CIRCUIT COURT OF THE {{ judicial_circuit }} JUDICIAL CIRCUIT", center=True, bold=True, space_after=0)
    _p(doc, "IN AND FOR {{ county|upper }} COUNTY, FLORIDA", center=True, bold=True, space_after=12)
    _p(doc, "STATE OF FLORIDA,", bold=True, space_after=0, double=False)
    _p(doc, "          Plaintiff,", space_after=0, double=False)
    _p(doc, "v.", space_after=0, double=False)
    _p(doc, "{{ defendant_name|upper }},", bold=True, space_after=0, double=False)
    _p(doc, "          Defendant.", space_after=6, double=False)
    _p(doc, "Case No. {{ case_number }}", bold=True, space_after=0, double=False)
    _p(doc, "Division: {{ court_division }}", space_after=12, double=False)
    _p(doc, "{{ motion_title|upper }}", center=True, bold=True, space_after=12)


def _signature(doc: Document) -> None:
    _p(doc, "Respectfully submitted this {{ filing_date }}.", space_before=12)
    _p(doc, "{{ counsel_firm }}", bold=True, space_before=18, space_after=0, double=False)
    for line in (
        "{{ counsel_name }}",
        "Florida Bar No. {{ counsel_bar }}",
        "{{ counsel_address }}",
        "Tel: {{ counsel_phone }}",
        "Email: {{ counsel_email }}",
        "Counsel for Defendant {{ defendant_name }}",
    ):
        _p(doc, line, space_after=0, double=False)
    _p(doc, "CERTIFICATE OF SERVICE", bold=True, space_before=18, first_indent=False)
    _p(doc, "{{ certificate_of_service }}", justify=True, first_indent=True)


def build_808() -> Path:
    """FORM 68710-8.08 — Right to Counsel."""
    doc = _base_doc()
    _caption(doc)

    _p(
        doc,
        "The Defendant, {{ defendant_name }}, by and through undersigned counsel, and pursuant "
        "to Rule 3.190, Florida Rules of Criminal Procedure, moves this Court to suppress the "
        "eyewitness identification evidence in this case. Specifically, the Defendant moves to "
        "suppress the following evidence:",
        justify=True,
        first_indent=True,
    )
    _p(
        doc,
        "1. The testimony of witness, {{ eyewitness_name }}, and Officer {{ officer_name }}, that on "
        "{{ identification_date }}, at {{ identification_place }}, {{ eyewitness_name }} identified the "
        "Defendant as the perpetrator of the {{ charged_offense }} charged in this case at a pretrial "
        "lineup in which the Defendant was not represented by counsel.",
        justify=True,
        first_indent=True,
    )
    _p(
        doc,
        "2. The testimony of {{ eyewitness_name }} at trial identifying the Defendant as the "
        "perpetrator of that {{ charged_offense }}.",
        justify=True,
        first_indent=True,
    )
    _p(
        doc,
        "The grounds for this Motion are that the extra-judicial identification of the Defendant by "
        "{{ eyewitness_name }} was made in violation of the Defendant’s right to counsel under the "
        "Sixth and Fourteenth Amendments to the United States Constitution and Article I, Section 16 "
        "of the Florida Constitution and Rules 3.130 and 3.111(a), Florida Rules of Criminal Procedure. "
        "These violations occurred because the Defendant was forced to participate in an extra-judicial "
        "identification procedure after the time that adversary judicial criminal proceedings had been "
        "initiated and without the presence of counsel.",
        justify=True,
        first_indent=True,
    )

    _p(doc, "FACTS", bold=True, space_before=12)
    _p(doc, "{{ factual_background }}", justify=True, first_indent=True)

    _p(doc, "ARGUMENT", bold=True, space_before=12)
    _p(doc, "{{ legal_argument }}", justify=True, first_indent=True)

    _p(doc, "WHEREFORE", bold=True, space_before=12)
    _p(
        doc,
        "WHEREFORE, the Defendant prays this Court will find that he or she was forced to participate "
        "in an extra-judicial identification procedure without the presence of counsel at a time when "
        "adversary judicial proceedings had been initiated and that the Court will, therefore, issue "
        "its Order suppressing {{ suppression_relief }}.",
        justify=True,
        first_indent=True,
    )
    _p(doc, "{{ prayer_for_relief }}", justify=True, first_indent=True)

    _signature(doc)

    out = OUT_DIR / "fl-8-08-suppress-id-right-to-counsel.docx"
    doc.save(str(out))
    return out


def build_809() -> Path:
    """FORM 68710-8.09 — Suggestive Identification Procedures."""
    doc = _base_doc()
    _caption(doc)

    _p(
        doc,
        "The Defendant, {{ defendant_name }}, by and through undersigned counsel, and pursuant "
        "to Rule 3.190, Florida Rules of Criminal Procedure, moves this Court to suppress the "
        "eyewitness identification evidence in this case. Specifically, the Defendant moves to "
        "suppress the following evidence:",
        justify=True,
        first_indent=True,
    )
    _p(
        doc,
        "1. The testimony of witnesses {{ eyewitness_name }} and Officer {{ officer_name }} that on "
        "{{ identification_date }}, at {{ identification_place }}, {{ eyewitness_name }} identified "
        "{{ identification_medium }} as the perpetrator of the "
        "{{ charged_offense }} charged in this case.",
        justify=True,
        first_indent=True,
    )
    _p(
        doc,
        "2. The testimony of {{ eyewitness_name }} at trial identifying the Defendant as the "
        "perpetrator of that {{ charged_offense }}.",
        justify=True,
        first_indent=True,
    )
    _p(
        doc,
        "The grounds for this Motion are that the extra-judicial procedure used on "
        "{{ identification_date }} was so unnecessarily suggestive as to give rise to a substantial "
        "and irreparable likelihood of misidentification in violation of the Fourteenth Amendment to "
        "the United States Constitution and Article I, Section 16 of the Florida Constitution.",
        justify=True,
        first_indent=True,
    )

    _p(doc, "FACTS", bold=True, space_before=12)
    _p(doc, "{{ factual_background }}", justify=True, first_indent=True)

    _p(doc, "ARGUMENT", bold=True, space_before=12)
    _p(doc, "{{ legal_argument }}", justify=True, first_indent=True)

    _p(doc, "WHEREFORE", bold=True, space_before=12)
    _p(
        doc,
        "WHEREFORE, the Defendant prays this Court will find that the "
        "{{ suggestive_procedure }} used by the police in this case was unnecessarily suggestive, "
        "and was conducive to a substantial and irreparable likelihood of misidentification and that "
        "the Court will, therefore, issue its Order suppressing {{ suppression_relief }}.",
        justify=True,
        first_indent=True,
    )
    _p(doc, "{{ prayer_for_relief }}", justify=True, first_indent=True)

    _signature(doc)

    out = OUT_DIR / "fl-8-09-suppress-id-suggestive.docx"
    doc.save(str(out))
    return out


def build_806_reference() -> Path:
    """Reference Word copy of jury instruction 3.9(c) (not a motion template)."""
    doc = _base_doc()
    _p(doc, "FLORIDA STANDARD JURY INSTRUCTION 3.9(c)", center=True, bold=True)
    _p(doc, "EYEWITNESS IDENTIFICATION", center=True, bold=True, space_after=12)
    _p(
        doc,
        "Source: LexisNexis Forms FORM 68710-8.06 — Florida Criminal Practice and Procedure. "
        "This file is stored for reference when preparing identification motions; it is not a "
        "fillable motion template.",
        italic=True,
        justify=True,
    )
    _p(
        doc,
        "Give if eyewitness identification is a disputed issue and if requested.",
        italic=True,
    )
    _p(
        doc,
        "You have heard testimony of eyewitness identification. In deciding how much weight to give "
        "to this testimony, you may consider the various factors mentioned in these instructions "
        "concerning credibility of witnesses.",
        justify=True,
        first_indent=True,
    )
    _p(
        doc,
        "In addition to those factors, in evaluating eyewitness identification testimony, you may "
        "also consider the capacity and opportunity to observe; whether the identification was the "
        "product of the witness’s own recollection or influence/suggestiveness; the circumstances "
        "of presentation for identification; inconsistent identifications; failures to identify; "
        "familiarity with the subject; lapses of time; cross-racial identification concerns; and the "
        "totality of the circumstances. Lineup administration requirements under Florida law "
        "(independent administrator or neutral alternatives, and required pre-lineup instructions) "
        "may also be considered in assessing reliability.",
        justify=True,
        first_indent=True,
    )
    _p(
        doc,
        "Full original text is retained in templates/lexis/sources/8.06-florida-jury-instruction-3.9c.pdf.",
        italic=True,
        space_before=12,
    )
    out = OUT_DIR / "fl-8-06-jury-instruction-eyewitness-id.docx"
    doc.save(str(out))
    return out


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for builder in (build_808, build_809, build_806_reference):
        path = builder()
        print(f"Wrote {path}")


if __name__ == "__main__":
    main()
